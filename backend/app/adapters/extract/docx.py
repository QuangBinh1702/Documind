"""Trích xuất DOCX — US-007 AC-3.

Chuẩn hoá về Markdown để giữ **cấp tiêu đề** và **danh sách**. Cấu trúc này
không chỉ để hiển thị: US-008 AC-3 ưu tiên cắt chunk tại ranh giới tiêu đề, và
`heading_path` của chunk được dựng từ đây.

DOCX không có khái niệm trang cố định — số trang phụ thuộc vào cách trình soạn
thảo dàn trang. Vì vậy toàn bộ tài liệu được coi là **một trang**, và trích dẫn
sẽ trỏ về vị trí ký tự thay vì số trang.
"""

from __future__ import annotations

from pathlib import Path

from app.adapters.extract.base import ExtractionError, ExtractResult, TextBuilder

__all__ = ["extract_docx"]

METHOD = "python-docx"

# Tên style của Word ứng với từng cấp tiêu đề, cả bản tiếng Anh và tiếng Việt.
_HEADING_PREFIXES = ("heading", "tiêu đề", "đầu đề")

_LIST_PREFIXES = ("list bullet", "list number", "list paragraph", "danh sách")


def _heading_level(style_name: str) -> int | None:
    """Lấy cấp tiêu đề từ tên style, ví dụ "Heading 2" → 2."""
    lowered = style_name.strip().lower()
    for prefix in _HEADING_PREFIXES:
        if lowered.startswith(prefix):
            tail = lowered[len(prefix) :].strip()
            if tail.isdigit():
                return min(int(tail), 6)
            return 1
    if lowered in {"title", "tiêu đề chính"}:
        return 1
    return None


def _is_list(style_name: str) -> bool:
    lowered = style_name.strip().lower()
    return any(lowered.startswith(p) for p in _LIST_PREFIXES)


def extract_docx(path: Path) -> ExtractResult:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError(
            "PYTHON_DOCX_MISSING",
            "Thiếu thư viện đọc DOCX trên máy chủ.",
        ) from exc

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise ExtractionError(
            "DOCX_UNREADABLE",
            "Không đọc được tệp DOCX — tệp có thể đã hỏng hoặc là định dạng .doc cũ.",
        ) from exc

    builder = TextBuilder()
    builder.start_page(1)

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name if para.style is not None else ""
        level = _heading_level(style)

        if level is not None:
            builder.add_block(f"{'#' * level} {text}")
        elif _is_list(style):
            builder.add_block(f"- {text}")
        else:
            builder.add_block(text)

    # Bảng: mỗi bảng thành một bảng Markdown để không mất quan hệ hàng/cột.
    for table in document.tables:
        rows = [
            [cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows
        ]
        rows = [r for r in rows if any(c for c in r)]
        if not rows:
            continue

        header, *body = rows
        lines = [
            "| " + " | ".join(header) + " |",
            "|" + "|".join("---" for _ in header) + "|",
        ]
        lines += ["| " + " | ".join(r) + " |" for r in body]
        builder.add_block("\n".join(lines))

    builder.end_page()
    return builder.build(method=METHOD)
